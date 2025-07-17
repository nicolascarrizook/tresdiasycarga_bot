#!/usr/bin/env python3
"""
Test script para procesar archivos DOCX
"""

import os
import sys
from pathlib import Path
from docx import Document

def test_docx_files():
    """Test para verificar que los archivos DOCX se pueden leer"""
    
    data_dir = Path("data/raw")
    
    if not data_dir.exists():
        print(f"❌ Error: Directorio {data_dir} no existe")
        return False
        
    docx_files = list(data_dir.glob("*.docx"))
    
    if not docx_files:
        print(f"❌ Error: No se encontraron archivos DOCX en {data_dir}")
        return False
    
    print(f"🔍 Archivos DOCX encontrados: {len(docx_files)}")
    
    for file_path in docx_files:
        try:
            print(f"\n📄 Procesando: {file_path.name}")
            
            # Intentar leer el documento
            doc = Document(str(file_path))
            
            # Contar párrafos
            paragraphs = len(doc.paragraphs)
            print(f"   📝 Párrafos: {paragraphs}")
            
            # Contar tablas
            tables = len(doc.tables)
            print(f"   📊 Tablas: {tables}")
            
            # Mostrar primera línea de texto no vacía
            first_text = None
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    first_text = paragraph.text.strip()[:100]
                    break
                    
            if first_text:
                print(f"   📖 Primer texto: {first_text}...")
            
            # Si hay tablas, mostrar información de la primera
            if tables > 0:
                table = doc.tables[0]
                rows = len(table.rows)
                cols = len(table.columns) if rows > 0 else 0
                print(f"   📋 Primera tabla: {rows} filas x {cols} columnas")
                
                # Mostrar primera fila
                if rows > 0:
                    first_row = []
                    for cell in table.rows[0].cells:
                        first_row.append(cell.text.strip())
                    print(f"   🔤 Primera fila: {first_row}")
                
        except Exception as e:
            print(f"   ❌ Error al procesar {file_path.name}: {e}")
            return False
    
    print(f"\n✅ Todos los archivos DOCX se pueden leer correctamente")
    return True

if __name__ == "__main__":
    success = test_docx_files()
    sys.exit(0 if success else 1)