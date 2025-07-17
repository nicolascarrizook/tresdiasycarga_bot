#!/usr/bin/env python3
"""
Procesador simple de datos DOCX para Sistema Mayra
"""

import os
import json
from pathlib import Path
from docx import Document
import sys

def process_almuerzos_cenas(file_path):
    """Procesar archivo de almuerzos y cenas"""
    doc = Document(file_path)
    recipes = []
    
    for table in doc.tables:
        if len(table.rows) > 1:  # Skip header
            for row in table.rows[1:]:  # Skip first row (header)
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0]:  # If has content
                    # Intentar identificar categorías
                    if any(cat in cells[0].upper() for cat in ['POLLO', 'CARNE', 'PESCADO', 'CERDO', 'VEGETAR', 'ENSALADA']):
                        category = cells[0]
                    else:
                        recipe = {
                            'name': cells[0],
                            'category': 'almuerzo_cena',
                            'type': 'principal',
                            'description': ' '.join(cells[1:]) if len(cells) > 1 else ''
                        }
                        recipes.append(recipe)
    
    return recipes

def process_desayunos(file_path):
    """Procesar archivo de desayunos"""
    doc = Document(file_path)
    recipes = []
    
    for table in doc.tables:
        if len(table.rows) > 1:
            # Las columnas son: Dulces | Salados | Colaciones
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                
                # Columna 0: Dulces
                if len(cells) > 0 and cells[0]:
                    recipes.append({
                        'name': cells[0],
                        'category': 'desayuno_merienda',
                        'type': 'dulce',
                        'description': ''
                    })
                
                # Columna 1: Salados
                if len(cells) > 1 and cells[1]:
                    recipes.append({
                        'name': cells[1],
                        'category': 'desayuno_merienda', 
                        'type': 'salado',
                        'description': ''
                    })
                
                # Columna 2: Colaciones
                if len(cells) > 2 and cells[2]:
                    recipes.append({
                        'name': cells[2],
                        'category': 'colacion',
                        'type': 'snack',
                        'description': ''
                    })
    
    return recipes

def process_all_docx():
    """Procesar todos los archivos DOCX"""
    data_dir = Path("data/raw")
    all_recipes = []
    
    # Procesar almuerzos/cenas
    almuerzos_file = data_dir / "almuerzoscena.docx"
    if almuerzos_file.exists():
        print("📍 Procesando almuerzos y cenas...")
        recipes = process_almuerzos_cenas(almuerzos_file)
        all_recipes.extend(recipes)
        print(f"   ✅ {len(recipes)} recetas encontradas")
    
    # Procesar desayunos
    desayunos_file = data_dir / "desayunos.docx"
    if desayunos_file.exists():
        print("📍 Procesando desayunos y meriendas...")
        recipes = process_desayunos(desayunos_file)
        all_recipes.extend(recipes)
        print(f"   ✅ {len(recipes)} recetas encontradas")
    
    # Guardar datos procesados
    output_dir = Path("data/processed")
    output_dir.mkdir(exist_ok=True)
    
    output_file = output_dir / "recipes.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_recipes, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ Total: {len(all_recipes)} recetas procesadas")
    print(f"📁 Guardadas en: {output_file}")
    
    # Mostrar algunas recetas de ejemplo
    print("\n📋 Ejemplos de recetas procesadas:")
    for i, recipe in enumerate(all_recipes[:5]):
        print(f"   {i+1}. {recipe['name']} ({recipe['category']} - {recipe['type']})")
    
    return all_recipes

if __name__ == "__main__":
    process_all_docx()