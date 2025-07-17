"""
Base parser for DOCX files with common functionality.
Provides foundation for all document parsers in the system.
"""

import os
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Union, Tuple
from pathlib import Path
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches

logger = logging.getLogger(__name__)


class BaseDocxParser(ABC):
    """
    Abstract base class for DOCX parsing with common functionality.
    Handles document loading, table extraction, and text processing.
    """
    
    def __init__(self, file_path: str):
        """Initialize parser with file path."""
        self.file_path = Path(file_path)
        self.document = None
        self.tables = []
        self.paragraphs = []
        self._validate_file()
        
    def _validate_file(self) -> None:
        """Validate that the file exists and is a DOCX file."""
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        if not self.file_path.suffix.lower() == '.docx':
            raise ValueError(f"File must be a DOCX file: {self.file_path}")
    
    def load_document(self) -> None:
        """Load the DOCX document and extract basic elements."""
        try:
            self.document = Document(self.file_path)
            self.tables = self.document.tables
            self.paragraphs = self.document.paragraphs
            logger.info(f"Loaded document: {self.file_path}")
            logger.info(f"Found {len(self.tables)} tables and {len(self.paragraphs)} paragraphs")
        except Exception as e:
            logger.error(f"Error loading document {self.file_path}: {str(e)}")
            raise
    
    def extract_tables_data(self) -> List[Dict[str, Any]]:
        """Extract data from all tables in the document."""
        tables_data = []
        
        for i, table in enumerate(self.tables):
            table_data = self._extract_table_data(table, i)
            if table_data:
                tables_data.append(table_data)
        
        return tables_data
    
    def _extract_table_data(self, table: Table, table_index: int) -> Dict[str, Any]:
        """Extract data from a single table."""
        try:
            rows = []
            headers = []
            
            # Extract headers from first row
            if table.rows:
                header_row = table.rows[0]
                headers = [self._clean_text(cell.text) for cell in header_row.cells]
            
            # Extract data rows
            for row_index, row in enumerate(table.rows[1:], start=1):
                row_data = []
                for cell in row.cells:
                    cell_text = self._clean_text(cell.text)
                    row_data.append(cell_text)
                
                if any(row_data):  # Only add non-empty rows
                    rows.append(row_data)
            
            return {
                'table_index': table_index,
                'headers': headers,
                'rows': rows,
                'row_count': len(rows),
                'column_count': len(headers)
            }
        
        except Exception as e:
            logger.error(f"Error extracting table {table_index}: {str(e)}")
            return {}
    
    def extract_paragraphs_data(self) -> List[Dict[str, Any]]:
        """Extract data from all paragraphs in the document."""
        paragraphs_data = []
        
        for i, paragraph in enumerate(self.paragraphs):
            paragraph_data = self._extract_paragraph_data(paragraph, i)
            if paragraph_data:
                paragraphs_data.append(paragraph_data)
        
        return paragraphs_data
    
    def _extract_paragraph_data(self, paragraph: Paragraph, index: int) -> Dict[str, Any]:
        """Extract data from a single paragraph."""
        try:
            text = self._clean_text(paragraph.text)
            
            if not text:
                return {}
            
            return {
                'index': index,
                'text': text,
                'style': paragraph.style.name if paragraph.style else None,
                'alignment': paragraph.alignment,
                'is_heading': self._is_heading(paragraph),
                'level': self._get_heading_level(paragraph)
            }
        
        except Exception as e:
            logger.error(f"Error extracting paragraph {index}: {str(e)}")
            return {}
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        if not text:
            return ""
        
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Handle common Spanish characters and formatting
        text = self._normalize_spanish_text(text)
        
        return text
    
    def _normalize_spanish_text(self, text: str) -> str:
        """Normalize Spanish text and handle common formatting issues."""
        # Replace common formatting issues
        replacements = {
            '…': '...',
            '–': '-',
            '—': '-',
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            '°': '°',
            '½': '1/2',
            '¼': '1/4',
            '¾': '3/4'
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
    
    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a heading."""
        if not paragraph.style:
            return False
        
        style_name = paragraph.style.name.lower()
        return 'heading' in style_name or 'title' in style_name
    
    def _get_heading_level(self, paragraph: Paragraph) -> Optional[int]:
        """Get heading level from paragraph style."""
        if not self._is_heading(paragraph):
            return None
        
        style_name = paragraph.style.name.lower()
        
        # Extract level from style name
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        
        if 'title' in style_name:
            return 1
        
        return None
    
    def find_tables_by_header(self, header_pattern: str) -> List[Tuple[int, Dict[str, Any]]]:
        """Find tables that contain a specific header pattern."""
        matching_tables = []
        tables_data = self.extract_tables_data()
        
        pattern = re.compile(header_pattern, re.IGNORECASE)
        
        for table_data in tables_data:
            headers = table_data.get('headers', [])
            if any(pattern.search(header) for header in headers):
                matching_tables.append((table_data['table_index'], table_data))
        
        return matching_tables
    
    def find_paragraphs_by_pattern(self, text_pattern: str) -> List[Dict[str, Any]]:
        """Find paragraphs that match a specific text pattern."""
        matching_paragraphs = []
        paragraphs_data = self.extract_paragraphs_data()
        
        pattern = re.compile(text_pattern, re.IGNORECASE)
        
        for paragraph_data in paragraphs_data:
            text = paragraph_data.get('text', '')
            if pattern.search(text):
                matching_paragraphs.append(paragraph_data)
        
        return matching_paragraphs
    
    def extract_nutritional_values(self, text: str) -> Dict[str, Union[float, None]]:
        """Extract nutritional values from text using regex patterns."""
        patterns = {
            'calories': r'(\d+(?:\.\d+)?)\s*(?:kcal|cal|calorías|calorias)',
            'protein': r'(\d+(?:\.\d+)?)\s*g?\s*(?:proteína|proteinas|protein)',
            'carbs': r'(\d+(?:\.\d+)?)\s*g?\s*(?:carbohidratos|carbs|hidratos)',
            'fat': r'(\d+(?:\.\d+)?)\s*g?\s*(?:grasa|grasas|fat)',
            'fiber': r'(\d+(?:\.\d+)?)\s*g?\s*(?:fibra|fiber)',
            'sugar': r'(\d+(?:\.\d+)?)\s*g?\s*(?:azúcar|azucar|sugar)',
            'sodium': r'(\d+(?:\.\d+)?)\s*(?:mg|g)?\s*(?:sodio|sodium)',
        }
        
        nutritional_values = {}
        
        for nutrient, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    value = float(match.group(1))
                    nutritional_values[nutrient] = value
                except ValueError:
                    nutritional_values[nutrient] = None
            else:
                nutritional_values[nutrient] = None
        
        return nutritional_values
    
    def extract_portions(self, text: str) -> List[Dict[str, Any]]:
        """Extract portion information from text."""
        # Pattern to match portions like "200g", "1 taza", "2 cucharadas"
        portion_patterns = [
            r'(\d+(?:\.\d+)?)\s*(g|gr|gramos?|kg|kilogramos?)',
            r'(\d+(?:\.\d+)?)\s*(ml|cc|litros?)',
            r'(\d+(?:\.\d+)?)\s*(taza|tazas|cup|cups)',
            r'(\d+(?:\.\d+)?)\s*(cucharada|cucharadas|cda|cdas|tbsp)',
            r'(\d+(?:\.\d+)?)\s*(cucharadita|cucharaditas|cdita|cditas|tsp)',
            r'(\d+(?:\.\d+)?)\s*(unidad|unidades|u|pieza|piezas)',
            r'(\d+(?:\.\d+)?)\s*(rodaja|rodajas|slice|slices)',
            r'(\d+(?:\.\d+)?)\s*(porción|porciones|portion|portions)'
        ]
        
        portions = []
        
        for pattern in portion_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    amount = float(match.group(1))
                    unit = match.group(2).lower()
                    portions.append({
                        'amount': amount,
                        'unit': unit,
                        'original_text': match.group(0)
                    })
                except (ValueError, IndexError):
                    continue
        
        return portions
    
    def extract_ingredients_list(self, text: str) -> List[str]:
        """Extract ingredients from text, handling common Spanish formats."""
        # Split by common delimiters
        ingredients = []
        
        # Try different splitting patterns
        if '•' in text:
            ingredients = text.split('•')
        elif '\n' in text:
            ingredients = text.split('\n')
        elif ',' in text:
            ingredients = text.split(',')
        elif ';' in text:
            ingredients = text.split(';')
        else:
            ingredients = [text]
        
        # Clean each ingredient
        cleaned_ingredients = []
        for ingredient in ingredients:
            ingredient = ingredient.strip()
            if ingredient and not ingredient.startswith('-'):
                # Remove bullet points and numbering
                ingredient = re.sub(r'^\d+\.?\s*', '', ingredient)
                ingredient = re.sub(r'^[-•*]\s*', '', ingredient)
                ingredient = ingredient.strip()
                if ingredient:
                    cleaned_ingredients.append(ingredient)
        
        return cleaned_ingredients
    
    @abstractmethod
    def parse(self) -> Dict[str, Any]:
        """Parse the document and return structured data."""
        pass
    
    @abstractmethod
    def validate_structure(self) -> bool:
        """Validate that the document has the expected structure."""
        pass
    
    def get_document_info(self) -> Dict[str, Any]:
        """Get basic information about the document."""
        if not self.document:
            self.load_document()
        
        return {
            'file_path': str(self.file_path),
            'file_size': self.file_path.stat().st_size,
            'table_count': len(self.tables),
            'paragraph_count': len(self.paragraphs),
            'has_tables': len(self.tables) > 0,
            'has_paragraphs': len(self.paragraphs) > 0
        }